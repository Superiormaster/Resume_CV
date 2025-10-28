import os
import io
import json
from datetime import datetime
import time
from shlex import split
from PIL.ImageOps import expand
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from weasyprint.layout.inline import split_text_box
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_login import LoginManager, login_user, login_required, current_user, UserMixin, logout_user
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, ForeignKey
from sqlalchemy.orm import scoped_session, sessionmaker, declarative_base
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import qrcode

# Flask App Configuration
BASE_DIR = os.path.abspath(os.path.dirname( __file__))
DB_PATH = os.path.join(BASE_DIR, 'data', 'app.db')
os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)

app = Flask(__name__)
app.config['SECRET_KEY'] = 'dev-secret-change-me'
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, "uploads")
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'docx', 'txt', 'pdf'}
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024

# Database Setup (SQLite)
engine = create_engine(f'sqlite:///{DB_PATH}', connect_args={'check_same_thread': False}, echo=False)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

db = SessionLocal()

# Database Models
class User(Base, UserMixin):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True)
    email = Column(String(200), unique=True, nullable=False)
    password_hash = Column(String(200), nullable=False)
    name = Column(String(200), nullable=True)

class Resume(Base):
    __tablename__ = "resumes"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    title = Column(String(200))
    template = Column(String(200))
    data_json = Column(Text)
    html_preview = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)

class UploadedFile(Base):
    __tablename__ = "uploads"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    filename = Column(String(255), nullable=False)
    filepath = Column(String(500), nullable=False)
    uploaded_at = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(engine)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

@login_manager.user_loader
def load_user(user_id):
    with SessionLocal() as db:
        return db.get(User, int(user_id))

# Routes
@app.route('/')
def index():
    """Homepage."""
    templates = os.listdir(os.path.join(BASE_DIR, 'templates', 'templates_variants'))
    return render_template('index.html', templates=templates)

# ---------------- Authentication ----------------

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Register a new user."""
    if request.method == 'POST':
        email = request.form['email'].strip().lower()
        password = request.form['password']
        name = request.form.get('name','').strip().lower()

        if not email or not password:
            flash("Email and password are required.")
            return redirect(url_for('register'))

        # Check whether url should be register or login
        with SessionLocal() as db:
            if db.query(User).filter_by(email=email).first():
                flash('Email already registered.')
                return redirect(url_for('login'))

        user = User(email=email, password_hash=generate_password_hash(password), name=name)
        db.add(user)
        db.commit()

        flash('Registered successfully. Please log in.')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login."""
    if request.method == 'POST':
        email = request.form['email'].lower().strip()
        password = request.form.get('password')

        if not email or not password:
            flash("Please enter both email and password.")
            return render_template(url_for('login'))

        with SessionLocal() as db:
            user = db.query(User).filter_by(email=email).first()

            if not user or not check_password_hash(user.password_hash, password):
                flash("Invalid Credentials")
                return render_template('login.html')

            # Log in the user
            login_user(user)
            flash('Login successful!')

            # Redirect to dashboard or the 'next' page if present
            next_page = request.args.get('next')
            return redirect(url_for('dashboard'))
    return render_template('login.html')

# ---------------- Dashboard ----------------

@app.route('/dashboard')
@login_required
def dashboard():
    """User dashboard to view saved resumes and uploaded files."""
    with SessionLocal() as db:
        resumes = (
            db.query(Resume)
            .filter_by(user_id=current_user.id)
            .order_by(Resume.created_at.desc())
            .all()
        )

        uploads = (
            db.query(UploadedFile)
            .filter_by(user_id=current_user.id)
            .order_by(UploadedFile.uploaded_at.desc())
            .all()
        )

    if not resumes:
        flash("You haven't created any resumes.", "info")

    return render_template('dashboard.html', resumes=resumes, uploads=uploads)

@app.route('/logout')
@login_required
def logout():
    """Logout user."""
    logout_user()
    flash("You have been logged out.", "info")
    return redirect(url_for('index'))

@app.route("/compose", methods=['GET', 'POST'])
@login_required
def compose():
    """Compose a new resume."""
    db = SessionLocal()
    try:
        if request.method == "POST":
            form = request.form.to_dict(flat=False)
            simple = {k: v[0] if isinstance(v, list) and len(v) == 1 else v for k, v in form.items()}
            structured = form_to_structured(simple)

            if not structured.get("full_name"):
                flash("Please enter your name")
                return redirect(url_for('compose'))

            html_preview = render_template("templates_variants/template_modern.html", data=structured, preview=True)

            resume = Resume(
                user_id=current_user.id,
                title=structured.get("full_name", "Untitled"),
                template=simple.get("template", "template_modern.html"),
                data_json=json.dumps(structured),
                html_preview=html_preview,
            )

            db.add(resume)
            db.commit()

            flash("Resume saved successfully!", "success")
            return redirect(url_for("dashboard"))

        templates_dir = os.path.join(BASE_DIR, 'templates', 'templates_variants')
        templates = os.listdir(templates_dir) if os.path.exists(templates_dir) else []
        return render_template('compose.html', templates=templates)

    finally:
        db.close()

# ---------------- Resume Preview ----------------

@app.route("/preview", methods=['POST'])
@login_required
def preview():

    form_data = request.form.to_dict(flat=False)
    simple = {k: v[0] if isinstance(v, list) and len(v) == 1 else v for k, v in form_data.items()}

    template =simple.get("template", 'template_modern.html')
    allowed_templates_dir = os.path.join(BASE_DIR, 'templates', 'templates_variants')
    allowed_templates = set(os.listdir(allowed_templates_dir)) if os.path.exists(allowed_templates_dir) else set()

    if template not in allowed_templates:
        flash("Invalid template selected.", "danger")
        return jsonify({"error": "Invalid template selected."}), 400

    structured = form_to_structured(simple)
    html = render_template(f"templates_variants/{template}", data = structured, preview=True)

    return jsonify({"html": html})

@app.route("/save", methods=['POST'])
@login_required
def save():
    """Save a completed resume to the database."""
    db = SessionLocal()
    try:
        form = request.form.to_dict(flat=False)
        simple = {k: v[0] if isinstance(v, list) and len(v) == 1 else v for k, v in form.items()}

        template = simple.get("template", "template_modern.html")
        templates_dir = os.path.join(BASE_DIR, 'templates', 'templates_variants')
        allowed_templates = set(os.listdir(templates_dir)) if os.path.exists(templates_dir) else set()

        if template not in allowed_templates:
            flash("Invalid template selected.", "danger")
            return redirect(url_for("compose"))

        structured = form_to_structured(simple)
        html = render_template(f"templates_variants/{template}", data=structured, preview=True)

        resume = Resume(
            user_id=current_user.id,
            title=structured.get('full_name','Untitled'),
            template=template,
            data_json=json.dumps(structured),
            html_preview=html,
        )
        db.add(resume)
        db.commit()
        flash("Resume saved successfully!", "success")

    except Exception as e:
        db.rollback()
        flash(f"Error saving resume: {e}", "danger")

    finally:
        db.close()

    return redirect(url_for('dashboard'))

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ---------------- File Upload ----------------

@app.route("/upload", methods=["POST"])
@login_required
def upload_file():
    """Upload a resume file."""
    if "file" not in request.files:
        flash("No file part.", "danger")
        return redirect(request.referrer or url_for("dashboard"))

    file = request.files["file"]

    if file.filename == "":
        flash("No file selected.", "warning")
        return redirect(request.referrer or url_for("dashboard"))

    if not allowed_file(file.filename):
        flash("Unsupported file type.", "danger")
        return redirect(request.referrer or url_for("dashboard"))

    filename = f"{int(time.time())}_{secure_filename(file.filename)}"
    save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(save_path)

    db = SessionLocal()
    try:
        uploaded = UploadedFile(
            user_id=current_user.id,
            filename=filename,
            filepath=save_path
        )
        db.add(uploaded)
        db.commit()
        flash("File uploaded successfully!", "success")
    except Exception as e:
        db.rollback()
        flash(f"Error saving upload: {e}", "danger")
    finally:
        db.close()

    return redirect(url_for("dashboard"))

@app.route("/uploads/download/<int:file_id>")
@login_required
def download_uploaded(file_id):
    db = SessionLocal()
    try:
        f = db.query(UploadedFile).filter_by(id=file_id, user_id=current_user.id)
        if not f:
            flash("File not found.", "danger")
            return redirect(request.referrer or url_for("dashboard"))

        return send_file(f.filepath, as_attachment=True, download_name=f.filename)
    except Exception as e:
        flash(f"Error downloading file: {e}", "danger")
        return redirect(request.referrer or url_for("dashboard"))
    finally:
        db.close()

@app.route("/uploads/delete/<int:file_id>")
@login_required
def delete_uploaded(file_id):
    db = SessionLocal()
    try:
        f = db.query(UploadedFile).filter_by(id=file_id, user_id=current_user.id).first()
        if not f:
            flash("File not found.", "danger")
            return redirect(request.referrer or url_for("dashboard"))

        if os.path.exists(f.filepath):
            os.remove(f.filepath)

        db.delete(f)
        db.commit()
        flash("File deleted successfully!", "success")
    except Exception as e:
        db.rollback()
        flash(f"Error deleting file: {e}", "danger")
    finally:
        db.close()

    return redirect(url_for("dashboard"))

@app.route("/resume/<int:resume_id>")
@login_required
def view_resume(resume_id):
    """Preview a specific resume."""
    db = SessionLocal()
    try:
        resume = db.query(Resume).filter_by(id=resume_id, user_id=current_user.id).first()
        if not resume:
            flash("Resume not found or you don't have access.", "warning")
            return redirect(url_for("dashboard"))

        templates_dir = os.path.join(BASE_DIR, 'templates', 'templates_variants')
        allowed_templates = set(os.listdir(templates_dir)) if os.path.exists(templates_dir) else set()
        if resume.template not in allowed_templates:
            flash("Template file missing or invalid.", "danger")
            return redirect(url_for("dashboard"))

        try:
            data = json.loads(resume.data_json)
        except json.JSONDecodeError:
            flash("Could not read resume.", "danger")
            return redirect(url_for("dashboard"))

        return render_template(f"templates_variants/{resume.template}", data=data, preview=False)

    except Exception as e:
        flash(f"Error loading resume: {e}", "danger")
        return redirect(url_for("dashboard"))

    finally:
        db.close()

def structured_to_plaintext(data):
    """Convert structured resume data to plain text."""
    return structured_to_text(data)

@app.route("/download/<int:resume_id>/<fmt>")
@login_required
def download_resume(resume_id, fmt):
    """Download a resume."""
    db = SessionLocal()
    print("Trying to download resume:", resume_id, fmt)
    try:
        r = db.query(Resume).filter_by(id=resume_id, user_id=current_user.id).first()
        if not r:
            flash("Resume not found or you don't have permission.", "danger")
            return redirect(url_for("dashboard"))

        data = json.loads(r.data_json)
        title = (data.get("full_name") or "resume").replace("_", " ")

        if fmt == "json":
            bio = io.BytesIO(json.dumps(data, indent=2).encode("utf-8"))
            bio.seek(0)
            return send_file(bio, as_attachment=True, download_name=f"{title}.json", mimetype="application/json")

        elif fmt == "txt":
            txt = structured_to_plaintext(data)
            bio = io.BytesIO(txt.encode("utf-8"))
            bio.seek(0)
            return send_file(bio, as_attachment=True, download_name=f"{title}.txt", mimetype="text/plain")

        elif fmt == "docx":
            docx_io = generate_docx(data)
            docx_io.seek(0)
            return send_file(docx_io, as_attachment=True, download_name=f"{title}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif fmt == "pdf":
            pdf_io = generate_pdf_reportlab(data)
            pdf_io.seek(0)
            return send_file(pdf_io, as_attachment=True, download_name=f"{title}.pdf", mimetype="application/pdf")

        else:
            flash(f"Unsupported download format.", "danger")
            return redirect(url_for("dashboard"))

    except Exception as e:
        print("Download error:", e)
        flash(f"Error generating resume: {e}", "danger")
        return redirect(url_for("dashboard"))
    finally:
        db.close()


def form_to_structured(d):
    structured = {}

    structured["full_name"] = d.get("full_name","").strip()
    structured["title"] = d.get("title","").strip()
    structured["email"] = d.get("email","").strip()
    structured["phone"] = d.get("phone","").strip()
    structured["summary"] = d.get("summary","").strip()

    skills = d.get("skills","").strip()
    # Normalize skills: allow list or comma-separated string
    if isinstance(skills, list):
        # if coming from JS as list
        structured["skills"] = [s.strip() for s in skills if s.strip()]
    else:
        #comma separated string
        structured["skills"] = [s.strip() for s in skills.split(",") if s.strip()]

    # Experience
    structured["experience"] = []
    for i in range(1,6):
        company = d.get(f"exp_company_{i}","").strip()
        role = d.get(f"exp_role_{i}","").strip()
        dates = d.get(f"exp_dates_{i}","").strip()
        desc = d.get(f"exp_desc_{i}","").strip()

        if company or role or desc:
            structured["experience"].append({"company":company, "role":role, "dates":dates, "desc":desc})

    # Education
    structured["education"] = []
    for i in range(1,4):
        school = d.get(f"edu_school_{i}","").strip()
        degree = d.get(f"edu_degree_{i}","").strip()
        years = d.get(f"edu_years_{i}","").strip()

        if school or degree:
            structured["education"].append({"school":school, "degree":degree, "years":years})

    # Single fields (linkedin, website)
    structured["linkedin"] = d.get(f"linkedin","").strip()
    structured["website"] = d.get(f"website","").strip()

    return structured

def structured_to_text(data):
    lines = []

    full_name = data.get("full_name","").strip()
    title = data.get("title","").strip()
    if full_name:
        lines.append(f"Full name: {full_name}")
    if title:
        lines.append(f"Title: {title}")
    lines.append("")

    #Contact
    lines.append("Contact:")
    if data.get("email"):
        lines.append(data.get("email").strip())
    if data.get("phone"):
        lines.append(data.get("phone").strip())
    if data.get("website"):
        lines.append(f"Website: {data['website'].strip()}")
    if data.get("linkedin"):
        lines.append(f"Linkedin: {data['linkedin'].strip()}")
    lines.append("")

    # Summary
    summary = (data.get("summary") or "").strip()
    if summary:
        lines.append("summary:")
        lines.append(data.get("summary","").strip())
        lines.append("")

    # Skills
    skills = data.get("skills", [])
    if skills:
        lines.append("Skills:")
        if isinstance(skills, list):
            lines.append(",".join(skills))
        else:
            lines.append(str(skills))
        lines.append("")

    # Experience
    experience = data.get("experience",[])
    if experience:
        lines.append("Experience:")
        for e in experience:
            role = e.get("role","")
            company = e.get("company","")
            dates = e.get("dates","")
            desc = e.get("desc","")
            line = f"{role} at {company}".strip()
            if dates:
                line += f" ({dates})"
            if line:
                lines.append(line)
            if desc:
                lines.append(desc)
            lines.append("")

    # Education
    education = data.get("education",[])
    if education:
        lines.append("Education:")
        for ed in education:
            degree = ed.get("degree","")
            school = ed.get("school","")
            years = ed.get("years","")
            line = f"{degree} - {school}".strip(" -")
            if years:
                line += f" ({years})"
            lines.append(line)
        lines.append("")

    return "\n".join(lines)

def generate_pdf_reportlab(data):
    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_io,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=30,
    )

    gold = colors.HexColor("#C9A227")   # Gold
    black = colors.HexColor("#0B0B0B")  # Black
    gray = colors.HexColor("#555555")

    # Custom styles
    styles = {
        "title": ParagraphStyle(
            "title",
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=black,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            fontName="Helvetica",
            fontSize=12,
            textColor=gold,
            spaceAfter=10,
        ),
        "section": ParagraphStyle(
            "section",
            fontName="Helvetica-Bold",
            fontSize=14,
            textColor=gold,
            spaceBefore=16,
            spaceAfter=6,
        ),
        "normal": ParagraphStyle(
            "normal",
            fontName="Helvetica",
            fontSize=11,
            textColor=black,
            leading=16,
        ),
    }

    story = []

    # Header
    story.append(Paragraph(data.get("full_name", "No Name"), styles["title"]))
    if data.get("title"):
        story.append(Paragraph(data.get("title", ""), styles["subtitle"]))
    story.append(Spacer(1, 10))

    # Contact Info
    contact = []
    if data.get("email"):
        contact.append(f"📧 {data['email']}")
    if data.get("phone"):
        contact.append(f"📞 {data['phone']}")
    if data.get("linkedin"):
        contact.append(f"🔗 {data['linkedin']}")
    if data.get("website"):
        contact.append(f"🌐 {data['website']}")
    story.append(Paragraph(" | ".join(contact), styles["normal"]))
    story.append(Spacer(1, 10))

    # Summary
    if data.get("summary"):
        story.append(Paragraph("Profile Summary", styles["section"]))
        story.append(Paragraph(data["summary"], styles["normal"]))

    # Skills
    skills = data.get("skills", [])
    if skills:
        story.append(Paragraph("Skills", styles["section"]))
        story.append(Paragraph(", ".join(skills), styles["normal"]))

    # Experience
    exp = data.get("experience", [])
    if exp:
        story.append(Paragraph("Experience", styles["section"]))
        for e in exp:
            story.append(
                Paragraph(f"<b>{e.get('role','')}</b> — {e.get('company','')} ({e.get('dates','')})", styles["normal"])
            )
            story.append(Paragraph(e.get("desc", ""), styles["normal"]))
            story.append(Spacer(1, 6))

    # Education
    edu = data.get("education", [])
    if edu:
        story.append(Paragraph("Education", styles["section"]))
        for ed in edu:
            story.append(
                Paragraph(f"{ed.get('degree','')} — {ed.get('school','')} ({ed.get('years','')})", styles["normal"])
            )
            story.append(Spacer(1, 6))

    # Optional QR Section
    if data.get("linkedin") or data.get("website"):
        qr_data = data.get("linkedin") or data.get("website")
        qr_img = qrcode.make(qr_data)
        bio = io.BytesIO()
        qr_img.save(bio, "PNG")
        bio.seek(0)
        story.append(Spacer(1, 12))
        story.append(Paragraph("Scan to view online:", styles["section"]))
        story.append(ImageReader(bio))

    doc.build(story)
    pdf_io.seek(0)
    return pdf_io

def split_text_to_lines(text, max_chars):
    """Split text into multiple lines without breaking words."""
    if not text:
        return []
    lines = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        cur = ""
        for w in words:
            if len(cur) + len(w) + (1 if cur else 0) <= max_chars:
                cur += (" " if cur else "") + w
            elif len(w) > max_chars:
                while len(w) > max_chars:
                    lines.append(w[:max_chars])
                    w = w[max_chars:]
                cur = w
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
    return lines

def generate_docx(data):
    doc = Document()

    doc.add_heading(data.get("full_name", ""), 0)
    if data.get("title"):
        doc.add_paragraph(data.get("title"))

    # Contact Info
    doc.add_paragraph(f"Email: {data.get('email','')}".strip())
    doc.add_paragraph(f"Phone: {data.get('phone','')}".strip())
    if data.get("website"):
        doc.add_paragraph(f"Website: {data.get('website')}")
    if data.get("linkedin"):
        doc.add_paragraph(f"Linkedin: {data.get('linkedin')}")

    # Profile
    doc.add_paragraph()
    doc.add_heading("Profile", level=1)
    doc.add_paragraph(data.get("summary", ""))

    # Skills
    doc.add_paragraph()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data.get("skills",[])))

    # Experience
    doc.add_paragraph()
    doc.add_heading("Experience", level=1)
    for e in data.get("experience", []):
        doc.add_paragraph(f"{e.get('role','')} - {e.get('company','')} | {e.get('dates','')}")
        doc.add_paragraph(e.get("desc",""))

    # Education
    doc.add_paragraph()
    doc.add_heading("Education", level=1)
    for ed in data.get("education",[]):
        doc.add_paragraph(f"{ed.get('degree','')} - {ed.get('school','')} ({ed.get('years','')})")

    # QR
    qr_data = data.get("linkedin") or data.get("website")
    if qr_data:
        qr_img = qrcode.make(qr_data)
        bio = io.BytesIO()
        qr_img.save(bio, "PNG")
        bio.seek(0)
        # Insert QR image (docx wants a file-like or path)
        doc.add_page_break()
        doc.add_heading("Scan this QR to view online:", level=1)
        doc.add_picture(bio, width = Inches(2))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

if __name__ == '__main__':
    app.run(debug=True)